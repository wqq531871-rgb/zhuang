# -*- coding: utf-8 -*-
"""Generate Word user manual for the packing visualization UI."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

out_dir = Path(__file__).resolve().parent
out_path = out_dir / "智能装箱规划系统_可视化界面使用说明.docx"

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_cn(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 12), bold=True)
    return p


def add_para(text, size=11, bold=False, space_after=6, first_line=True):
    p = doc.add_paragraph()
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    r0 = p.add_run("提示：")
    set_run_font(r0, bold=True, color=RGBColor(0x1D, 0x4E, 0xD8))
    r1 = p.add_run(text)
    set_run_font(r1, color=RGBColor(0x33, 0x41, 0x55))
    return p


# ---- Cover ----
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("面向控序混码场景")
set_run_font(r, size=22, bold=True)

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title2.add_run("智能装箱规划系统")
set_run_font(r, size=26, bold=True)

title3 = doc.add_paragraph()
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title3.add_run("可视化界面使用说明")
set_run_font(r, size=18, bold=True, color=RGBColor(0x25, 0x63, 0xEB))

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("主界面：Industrial Packing Workbench V3 Clean")
set_run_font(r, size=11, color=RGBColor(0x64, 0x74, 0x8B))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("一键装箱 · 结果分析 · 托盘切换 · 稳定性评估")
set_run_font(r, size=11, color=RGBColor(0x64, 0x74, 0x8B))

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "文档版本：V1.0\n"
    "适用对象：现场操作人员、工艺/计划人员、系统联调人员\n"
    "说明范围：界面操作与结果查看（不含算法原理）"
)
set_run_font(r, size=10.5, color=RGBColor(0x47, 0x55, 0x69))

doc.add_page_break()

# ---- TOC ----
add_heading_cn("目录", 1)
for item in [
    "1  系统简介",
    "2  启动与关闭",
    "3  界面总览",
    "4  日常主流程：Excel 单次装箱",
    "5  接口模式：从 WCS 拉取库存并装箱",
    "6  查看与分析装箱结果",
    "7  下传 WCS 与现场码垛",
    "8  历史结果与日志",
    "9  高级功能（一般不用）",
    "10 常见问题与排查",
]:
    add_para(item, first_line=False, space_after=4)

doc.add_page_break()

# ---- 1 ----
add_heading_cn("1  系统简介", 1)
add_para(
    "本系统通过可视化工作台完成装箱计算、三维结果查看、稳定性复核，以及与现场 WCS/码垛链路的衔接。"
    "日常使用以顶栏操作为主：选择数据来源 → 一键装箱 → 在中间工作区查看托盘与箱子 → 需要时下传到现场。"
)
add_para(
    "本文档只介绍界面怎么用，不解释装箱算法内部细节。打开软件后，右上角状态胶囊会显示当前状态："
    "空闲 / 运行中 / 已完成 / 失败。"
)

add_heading_cn("1.1  你能用它做什么", 2)
for t in [
    "用 Excel 订单数据做一次装箱计算，并自动加载结果。",
    "按设定周期从 WCS 接口拉取库存数据，持续或单次计算。",
    "在三维视图中查看托盘码放效果，逐步回放装箱过程。",
    "筛选成功/失败托盘，查看箱子列表、失败原因与稳定性指标。",
    "把达标托盘整盘下传到 WCS；在现场码垛面板跟踪自动下传与应急补发。",
    "打开独立三维演示窗口，配合现场选定托盘做可视化演示。",
]:
    add_bullet(t)

add_heading_cn("1.2  启动时自动做了什么", 2)
add_para(
    "启动主界面时，系统会在后台自动拉起局域网接收端（用于对接现场接口）。"
    "关闭主窗口时会自动停止该接收端。接收端根地址与 Swagger 页面地址会写在底部日志中，便于联调查看。"
)

# ---- 2 ----
add_heading_cn("2  启动与关闭", 1)
add_heading_cn("2.1  推荐启动方式（Windows）", 2)
add_para("双击或运行以下启动脚本：", first_line=False)
add_bullet(
    "packing-system\\tools\\windows\\start_realtime_dashboard_v3_clean.bat",
    bold_prefix="脚本路径：",
)
add_para("也可在 packing-system 目录下用命令启动：", first_line=False)
add_bullet("python ui/realtime_dashboard_v3_clean.py --project .", bold_prefix="命令：")

add_heading_cn("2.2  关闭", 2)
add_para(
    "直接关闭主窗口即可。系统会停止后台装箱进程（若仍在运行）以及自动启动的局域网接收端。"
    "若三维演示窗口已单独打开，请一并关闭。"
)
add_note(
    "首次使用请确认本机已安装 Python，并能正常打开图形界面；"
    "若三维视图提示缺少依赖，请联系维护人员安装相关组件。"
)

# ---- 3 ----
add_heading_cn("3  界面总览", 1)
add_para(
    "主界面大致分为四个区域：顶栏、左侧流程区、中间工作区、右侧托盘摘要；底部为可展开的运行日志。"
)

add_heading_cn("3.1  顶栏（主操作入口）", 2)
add_bullet("显示空闲 / 运行中 / 已完成 / 失败。", bold_prefix="状态胶囊：")
add_bullet(
    "接口持续运行 / 接口单次运行 / Excel 单次运行 / 接口运行至成功。",
    bold_prefix="运行方式：",
)
add_bullet(
    "仅接口类模式可用，单位为秒，表示拉取库存并计算的周期。",
    bold_prefix="拉取间隔：",
)
add_bullet("选择装箱输入 Excel（Excel 模式）。", bold_prefix="选择 Excel：")
add_bullet("按当前运行方式启动装箱。", bold_prefix="一键装箱：")
add_bullet("停止当前正在运行的装箱任务。", bold_prefix="停止：")
add_bullet("手动打开历史 JSON 结果文件。", bold_prefix="打开结果文件：")
add_bullet(
    "「当前」= 最近一次结果；其余为最近约 50 条历史记录。",
    bold_prefix="结果历史下拉框：",
)
add_bullet(
    "切换算法目录、配置文件、查看当前设置、按配置复跑（日常一般不需要）。",
    bold_prefix="算法设置：",
)

add_heading_cn("3.2  左侧流程区", 2)
add_bullet("进度条与文字提示当前计算状态。", bold_prefix="① 运行状态：")
add_bullet(
    "按状态（全部/SUCCESS/FAILED/UNKNOWN）、订单、关键词筛选托盘；支持分页浏览。",
    bold_prefix="② 筛选托盘：",
)
add_bullet(
    "总托盘数、成功/失败托盘、平均填充率等汇总指标。",
    bold_prefix="概览指标卡：",
)
add_bullet(
    "勾选一个或多个达标且未下传的托盘，整盘下传到 WCS。",
    bold_prefix="↓ 下传 WCS：",
)
add_bullet(
    "显示当前选定托盘的订单、进度、是否旋转、码放队列；可刷新、打开三维演示、应急补发。",
    bold_prefix="▣ 现场码垛：",
)
add_bullet(
    "默认收起。只影响界面上的稳定性复核评价，不会改动已算好的装箱结果。",
    bold_prefix="③ 高级参数：",
)

add_heading_cn("3.3  中间工作区（四个页签）", 2)
add_bullet(
    "每页最多显示 6 个托盘三维总览；可拖动/缩放/旋转视角；点击托盘联动右侧信息；"
    "可进入单托盘放大与动画回放。",
    bold_prefix="3D装箱视图：",
)
add_bullet(
    "按装箱执行顺序列出当前托盘内的箱子；可上移/下移后点「更新」写回顺序。",
    bold_prefix="箱子列表：",
)
add_bullet(
    "展示失败箱或失败托盘及原因；双击可跳到对应托盘。",
    bold_prefix="失败列表：",
)
add_bullet(
    "综合评分、稳定等级、平均支撑率、风险箱数量，以及分项指标表。",
    bold_prefix="稳定性分析：",
)

add_heading_cn("3.4  右侧摘要与底部日志", 2)
add_para(
    "右侧展示当前选中托盘的关键信息与关键指标（填充率、指数、综合评分、稳定等级、箱数、重量、"
    "高度利用率、支撑率、重心偏移等），并给出用人话描述的操作建议。"
    "底部可「展开日志 / 收起日志」，查看运行过程与告警；支持「清空」。"
)

# ---- 4 ----
add_heading_cn("4  日常主流程：Excel 单次装箱", 1)
add_para(
    "这是最常用的离线/试算流程。适合已有订单 Excel、想快速看装箱效果的场景。"
)

add_heading_cn("4.1  操作步骤", 2)
add_numbered(
    "在顶栏「运行方式」中选择「Excel 单次运行」。此时「拉取间隔」会变为不可用。"
)
add_numbered("点击「选择 Excel」，在文件对话框中选择 .xlsx / .xls 输入文件。")
add_numbered(
    "系统会自动把 Excel 复制到工作区数据目录，并生成本次临时运行配置；相关信息写入底部日志。"
)
add_numbered("点击「一键装箱」。若尚未选过 Excel，系统会先弹出选择框。")
add_numbered(
    "等待状态变为「已完成」。结果会自动加载到界面，无需再手动找结果文件。"
)
add_numbered(
    "在左侧筛选托盘，在「3D装箱视图」中查看码放效果，必要时查看「失败列表」与「稳定性分析」。"
)

add_heading_cn("4.2  选择 Excel 时系统帮你做了什么", 2)
add_bullet("自动复制到项目数据目录，减少中文路径、空格路径带来的问题。")
add_bullet("自动识别工作表，并生成本次临时配置；无需手工改配置文件。")
add_bullet("若工作表不完整，日志中会出现「警告」提示，请按提示检查 Excel 内容。")
add_note(
    "Excel 模式下，一次「一键装箱」只算一轮；算完会自动进入结果展示。"
    "若要换一份 Excel，重新「选择 Excel」后再点「一键装箱」。"
)

# ---- 5 ----
add_heading_cn("5  接口模式：从 WCS 拉取库存并装箱", 1)
add_para(
    "当现场库存由 WCS 接口提供时，使用接口类运行方式。"
    "顶栏「拉取间隔」用于控制多久拉一次数据并计算一次。"
)

add_heading_cn("5.1  四种运行方式怎么选", 2)
add_bullet(
    "按拉取间隔反复拉取并计算，直到你点击「停止」。适合持续盯库存、持续出方案。",
    bold_prefix="接口持续运行：",
)
add_bullet(
    "只拉取并计算一次，然后自动结束。适合临时试一次接口数据。",
    bold_prefix="接口单次运行：",
)
add_bullet(
    "按间隔反复计算，一旦出现成功（达标）托盘就自动停止。适合「等到有可用方案就停」。",
    bold_prefix="接口运行至成功：",
)
add_bullet("不拉接口，只用本地 Excel（见第 4 章）。", bold_prefix="Excel 单次运行：")

add_heading_cn("5.2  操作步骤", 2)
add_numbered("选择「接口持续运行 / 接口单次运行 / 接口运行至成功」之一。")
add_numbered("设置「拉取间隔」（例如 200 秒；允许范围 1–86400 秒）。")
add_numbered("点击「一键装箱」。系统会按所选方式启动接口装箱服务。")
add_numbered("运行过程中，新产生的有效结果会自动推送到界面；历史下拉框也会更新。")
add_numbered(
    "持续运行时，需要结束请点「停止」。单次或「至成功」模式会在条件满足后自行结束。"
)
add_note(
    "接口模式依赖数据库与接口配置（由维护人员预先配置）。"
    "界面日常只改运行方式与拉取间隔，一般不要改「算法设置」。"
)

# ---- 6 ----
add_heading_cn("6  查看与分析装箱结果", 1)

add_heading_cn("6.1  筛选与切换托盘", 2)
add_para("左侧「筛选托盘」区域提供：", first_line=False)
add_bullet("全部 / SUCCESS（达标）/ FAILED（未达标）/ UNKNOWN。", bold_prefix="状态：")
add_bullet("按订单过滤。", bold_prefix="订单：")
add_bullet("支持托盘 ID、箱型、订单号关键词。", bold_prefix="搜索：")
add_bullet("上一页 / 下一页；中间三维区每页最多 6 个托盘。", bold_prefix="分页：")
add_para("点击某个托盘卡片后，右侧摘要与各页签数据会联动更新到该托盘。")

add_heading_cn("6.2  三维总览与单托盘放大", 2)
add_para("在「3D装箱视图」中：", first_line=False)
add_bullet("鼠标拖动可旋转视角，滚轮缩放；拖动与「点击选中」已区分，拖动时不会误切换托盘。")
add_bullet("点击托盘卡片上的放大按钮，进入单托盘大视图。")
add_bullet(
    "大视图工具栏：最终 / 播放 / 暂停 / 前一步 / 后一步 / 重置，用于逐步观察装箱过程。"
)
add_bullet(
    "「速度」滑条可调动画快慢；「着色」可按支撑风险、重量、层高、箱型、箱子区分等切换。"
)
add_bullet("勾选「吸盘」「风险箱」可叠加显示相关提示。")
add_bullet("点「返回总览」回到六托盘网格。")

add_heading_cn("6.3  箱子列表", 2)
add_para(
    "按执行顺序排列当前托盘内的箱子。选中一行后，可用「上移」「下移」调整顺序，"
    "再点「更新」写回本次计算结果相关文件。调整顺序属于现场微调能力，请确认后再更新。"
)

add_heading_cn("6.4  失败列表", 2)
add_para(
    "优先显示失败箱信息；若结果中没有逐箱失败字段，则汇总显示失败托盘、缺口与低填充率等原因。"
    "双击某一行可跳转到对应托盘，便于快速定位问题。"
)

add_heading_cn("6.5  稳定性分析", 2)
add_para(
    "顶部四个大指标：综合评分、稳定等级、平均支撑率、风险箱数量。"
    "下方表格给出分项分数、当前值、状态与说明。右侧「操作建议」会用通俗语言提示主要风险。"
)
add_para(
    "左侧「高级参数」中的参数方案（标准方案 / 保守方案 / 自定义）只影响前端稳定性复核观感，"
    "不会重新计算装箱布局。日常保持「标准方案」即可；需要更严格排查风险时可切换「保守方案」。"
)

# ---- 7 ----
add_heading_cn("7  下传 WCS 与现场码垛", 1)

add_heading_cn("7.1  下传 WCS（整盘下传）", 2)
add_para(
    "当数据库中存在「达标且尚未下传」的托盘时，左侧「下传 WCS」区域会提示未下传数量。"
)
add_numbered("点击「选择托盘下传…」。")
add_numbered("在弹窗中勾选一个或多个托盘，确认下传。")
add_numbered(
    "成功后会提示下传托盘数，并在库中标记为已下传；失败时请查看底部日志。"
)
add_note(
    "若提示「暂无未下传托盘可下传」，说明当前没有可下传的达标托盘，或数据库暂不可用。"
)

add_heading_cn("7.2  现场码垛面板", 2)
add_para(
    "该面板跟踪当前由 WCS 选定的托盘，并显示码放队列。"
    "正常情况下箱子到达且姿态就绪后会自动下传；面板约每 2 秒自动刷新，也可手动点「刷新」。"
)
add_bullet("当前订单号。", bold_prefix="订单：")
add_bullet("第几箱 / 总箱数。", bold_prefix="进度：")
add_bullet("当前箱是否需要旋转 90°。", bold_prefix="是否旋转：")
add_bullet("排队中 / 已下传 / 失败等状态说明。", bold_prefix="状态与队列：")

add_heading_cn("7.3  打开三维演示", 2)
add_para(
    "点击「打开三维演示」会以独立窗口启动机器人三维仿真界面，"
    "按当前选定托盘从数据库加载整盘方案。适合现场讲解或联调演示。"
    "关闭主界面时如三维窗口仍开着，请手动关闭。"
)

add_heading_cn("7.4  应急补发", 2)
add_para(
    "仅在自动下传卡住时使用。在队列中选中状态为「排队中」、且符合当前应下传序号的箱子，"
    "点击「应急补发」，确认后手动补发这一箱。系统强制按顺序下传，不能跳号。"
)

# ---- 8 ----
add_heading_cn("8  历史结果与日志", 1)
add_heading_cn("8.1  结果历史下拉框", 2)
add_para(
    "顶栏下拉框中，「当前」表示最新一次装箱结果；其下为历史记录（约最近 50 条），"
    "标签中通常含时间与来源（如 Excel / 接口 / 达标等）。"
    "切换后界面会加载对应结果，并跳到「3D装箱视图」。"
)

add_heading_cn("8.2  打开结果文件", 2)
add_para(
    "若需要查看更早、或不在历史列表中的 JSON 结果，可点「打开结果文件」手动选择。"
    "文件需为有效装箱结果（包含托盘列表）。"
)

add_heading_cn("8.3  底部日志", 2)
add_para(
    "日志记录选择 Excel、生成配置、后端运行、下传 WCS、现场码垛、接收端启停等信息。"
    "出现异常时，优先展开日志查看报错行，便于反馈给维护人员。"
)

# ---- 9 ----
add_heading_cn("9  高级功能（一般不用）", 1)
add_para("顶栏「算法设置」菜单面向维护与联调，日常操作通常无需使用：")
add_bullet("切换装箱工程根目录。", bold_prefix="选择算法目录…：")
add_bullet("指定某份 YAML 配置。", bold_prefix="选择配置文件…：")
add_bullet(
    "弹窗显示当前工程目录、配置路径、本次输出路径等。",
    bold_prefix="查看当前设置：",
)
add_bullet(
    "不走「选择 Excel」流程，直接按当前配置再跑一次。",
    bold_prefix="按当前配置复跑算法：",
)
add_note(
    "日常标准路径仍是：选择运行方式 →（Excel 模式选文件）→ 一键装箱。"
    "只有更换工程或参数文件时才需要改这里。"
)

# ---- 10 ----
add_heading_cn("10  常见问题与排查", 1)

add_heading_cn("10.1  点了「一键装箱」没反应或立刻失败", 2)
add_bullet("Excel 模式是否已选到有效的 .xlsx/.xls 文件。")
add_bullet("接口模式时数据库/接口是否可用（看日志中的错误信息）。")
add_bullet(
    "是否已有任务在跑（提示「后端装箱正在运行」时，先「停止」或等其结束）。"
)

add_heading_cn("10.2  三维区域空白或提示缺少 3D 依赖", 2)
add_para(
    "说明本机三维显示组件未就绪。可先用「箱子列表 / 失败列表 / 稳定性分析」查看结果，"
    "并联系维护人员安装三维依赖。"
)

add_heading_cn("10.3  局域网接收端启动失败", 2)
add_para(
    "日志可能提示端口被占用（常见为 8093）或找不到接收端配置。"
    "请确认端口未被其他程序占用，或联系维护人员检查 local_wcs_receiver 配置。"
)

add_heading_cn("10.4  下传 WCS 失败", 2)
add_bullet("确认库中确有未下传达标托盘。")
add_bullet("查看日志中的下传 URL 与返回信息。")
add_bullet("网络或对方接口异常时，稍后重试或联系联调人员。")

add_heading_cn("10.5  现场码垛无法应急补发", 2)
add_bullet("是否选中了「排队中」的箱子。")
add_bullet("是否按顺序：只能补发当前应下传的那一箱，不能跳号。")
add_bullet("是否已选定托盘（面板提示「等待 WCS 选定托盘」时，需先完成选定）。")

add_heading_cn("10.6  想重新看刚才的结果", 2)
add_para(
    "在顶栏结果历史下拉框选「当前」或对应历史记录；或用「打开结果文件」手动加载。"
)

# ---- Appendix ----
add_heading_cn("附录 A  推荐日常操作速查", 1)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "场景"
hdr[1].text = "怎么做"
hdr[2].text = "看哪里"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, bold=True, size=10)

rows = [
    (
        "用 Excel 试算一次",
        "运行方式选 Excel → 选择 Excel → 一键装箱",
        "3D装箱视图 + 右侧摘要",
    ),
    (
        "接口持续盯库存",
        "选接口持续运行 → 设拉取间隔 → 一键装箱；结束点停止",
        "结果历史自动更新",
    ),
    (
        "等到有达标托盘就停",
        "选接口运行至成功 → 一键装箱",
        "成功后自动结束并加载",
    ),
    (
        "查某个失败托盘",
        "左侧状态选 FAILED，或打开失败列表双击",
        "失败列表 / 三维视图",
    ),
    ("整盘下传现场", "下传 WCS → 勾选托盘 → 确认", "弹窗成功提示 + 日志"),
    (
        "现场卡住补一箱",
        "现场码垛队列选中排队中箱子 → 应急补发",
        "现场码垛面板状态",
    ),
    ("演示三维码垛", "选定托盘后点打开三维演示", "独立三维窗口"),
]
for a, b, c in rows:
    row = table.add_row().cells
    row[0].text = a
    row[1].text = b
    row[2].text = c
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10)

doc.add_paragraph()
add_heading_cn("附录 B  界面名词对照", 1)
add_bullet("顶栏胶囊显示的空闲/运行中/已完成/失败。", bold_prefix="状态：")
add_bullet(
    "计算成功、可进入后续下传等流程的托盘（界面中常对应 SUCCESS）。",
    bold_prefix="达标托盘：",
)
add_bullet(
    "未满足要求的托盘（界面中常对应 FAILED）。",
    bold_prefix="未达标托盘：",
)
add_bullet("把已达标托盘的方案整盘推送给现场 WCS。", bold_prefix="下传：")
add_bullet(
    "箱子在托盘上的执行序号，可视化播放与现场下传都按此顺序。",
    bold_prefix="执行序（seq）：",
)
add_bullet(
    "启动主界面时自动后台启动的局域网服务，用于对接现场接口。",
    bold_prefix="局域网接收端：",
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("— 文档结束 —")
set_run_font(r, size=10, color=RGBColor(0x94, 0xA3, 0xB8))

footer_note = doc.add_paragraph()
footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_note.add_run("如界面按钮文案与本文略有差异，以实际软件界面为准。")
set_run_font(r, size=9, color=RGBColor(0x94, 0xA3, 0xB8))

doc.save(out_path)
print(str(out_path))
print("OK")
